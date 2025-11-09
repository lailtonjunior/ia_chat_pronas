"""
Serviço de geração de relatórios (PDF/DOCX)
"""

from __future__ import annotations

import logging
import os
import tempfile
import uuid
from datetime import datetime, timedelta
from pathlib import Path

import boto3
from docx import Document as DocxDocument
from weasyprint import HTML
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select

from app.config import settings
from app.db.database import AsyncSessionLocal
from app.models.generated_report import GeneratedReport, ReportFormat, ReportStatus
from app.models.project import Project
from app.services.notification_service import NotificationService
from app.models.notification import NotificationType, NotificationSeverity, NotificationChannel

logger = logging.getLogger(__name__)


class ReportService:
    STORAGE_DIR = Path(settings.UPLOAD_PATH).parent / "reports"

    @staticmethod
    async def request_report(
        db: AsyncSession,
        project_id: uuid.UUID,
        report_format: ReportFormat,
    ) -> GeneratedReport:
        report = GeneratedReport(project_id=project_id, format=report_format)
        db.add(report)
        await db.commit()
        await db.refresh(report)
        return report

    @staticmethod
    async def generate(report_id: uuid.UUID, report_format: ReportFormat) -> None:
        async with AsyncSessionLocal() as session:
            result = await session.execute(
                select(GeneratedReport).where(GeneratedReport.id == report_id)
            )
            report = result.scalar_one_or_none()
            if not report:
                logger.error("Report %s não encontrado", report_id)
                return

            report.status = ReportStatus.PROCESSING
            await session.commit()

            project = await session.get(Project, report.project_id)
            if not project:
                report.status = ReportStatus.FAILED
                await session.commit()
                return

            file_path = await ReportService._generate_file(project, report_format)
            file_url = await ReportService._upload_file(file_path)

            report.status = ReportStatus.READY
            report.file_url = file_url
            report.expires_at = datetime.utcnow() + timedelta(
                minutes=settings.REPORT_URL_EXPIRATION_MINUTES
            )
            await session.commit()

            await NotificationService.create_notification(
                session,
                user_id=project.user_id,
                title="Relatório pronto",
                message=f"O relatório do projeto \"{project.title}\" está disponível.",
                notification_type=NotificationType.REPORT_READY,
                severity=NotificationSeverity.SUCCESS,
                channel=NotificationChannel.IN_APP,
                data={"project_id": str(project.id), "report_id": str(report.id), "url": report.file_url},
            )

    @staticmethod
    async def _generate_file(project: Project, report_format: ReportFormat) -> Path:
        ReportService.STORAGE_DIR.mkdir(parents=True, exist_ok=True)
        filename = f"{project.id}_{datetime.utcnow().strftime('%Y%m%d%H%M%S')}.{report_format.value}"
        output_path = ReportService.STORAGE_DIR / filename

        if report_format == ReportFormat.PDF:
            html = f"""
            <html>
                <head><meta charset='utf-8'><style>body{{font-family: Arial;}}</style></head>
                <body>
                    <h1>{project.title}</h1>
                    <p>{project.description or ''}</p>
                </body>
            </html>
            """
            HTML(string=html).write_pdf(str(output_path))
        else:
            doc = DocxDocument()
            doc.add_heading(project.title, 0)
            doc.add_paragraph(project.description or "")
            doc.save(str(output_path))

        return output_path

    @staticmethod
    async def _upload_file(file_path: Path) -> str:
        if not settings.REPORT_STORAGE_BUCKET:
            return str(file_path)

        s3 = boto3.client(
            "s3",
            aws_access_key_id=settings.AWS_ACCESS_KEY_ID,
            aws_secret_access_key=settings.AWS_SECRET_ACCESS_KEY,
            region_name=settings.AWS_REGION_NAME,
        )
        key = f"reports/{file_path.name}"
        s3.upload_file(str(file_path), settings.REPORT_STORAGE_BUCKET, key)
        return f"s3://{settings.REPORT_STORAGE_BUCKET}/{key}"
